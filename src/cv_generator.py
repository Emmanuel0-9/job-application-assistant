import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from src.models import CV
from config import OUTPUT_DIR

# Paleta
BLUE  = RGBColor(46, 116, 181)
DARK  = RGBColor(32, 32, 32)
GRAY  = RGBColor(102, 102, 102)


def _run(para, text: str, size: int = 10, bold: bool = False,
         color: RGBColor = None, font: str = "Calibri"):
    run = para.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color if color else DARK
    return run


def _rule(para):
    """Línea horizontal azul bajo el párrafo."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E74B5")
    pBdr.append(bot)
    pPr.append(pBdr)


def _section(doc, title: str):
    para = doc.add_paragraph()
    _run(para, title, size=12, bold=True, color=BLUE)
    _rule(para)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)


# Nombres de dispositivo de Windows. Escribir en ellos no crea un archivo: habla
# con el dispositivo. Se comprobo en Windows 11 que "NUL" a secas se acepta y deja
# el archivo en 0 bytes — el CV se perderia en silencio.
_RESERVADOS_WINDOWS = {
    "CON", "PRN", "AUX", "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


def _nombre_archivo_seguro(nombre: str) -> str:
    """Reduce un nombre de archivo a algo que no pueda escapar de OUTPUT_DIR.

    El nombre se arma con datos que vienen de afuera (la empresa de la oferta,
    procesada por la IA). Sin esto, una empresa llamada "../../.." haria que
    doc.save() escribiera fuera de la carpeta output/. Se valida aqui, en la
    frontera: se descarta todo lo que no sea una letra, numero, guion o punto.
    """
    base = Path(nombre).name                     # descarta cualquier ruta: ../ , \ , /
    base = re.sub(r"[^A-Za-z0-9._-]", "_", base)  # solo caracteres seguros
    base = base.lstrip(".")                       # sin punto inicial (archivos ocultos)
    if not base:
        return "cv.docx"
    # Si la raiz es un dispositivo de Windows, se le antepone algo para desactivarlo.
    if base.split(".")[0].upper() in _RESERVADOS_WINDOWS:
        base = f"cv_{base}"
    return base


def generate_cv_docx(cv: CV, filename: str) -> Path:
    """Genera un .docx ATS-friendly y devuelve la ruta."""
    doc = Document()

    # Márgenes 2 cm
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    # ── ENCABEZADO ────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, cv.personal.name.upper(), size=20, bold=True, color=BLUE)

    contact = " | ".join(filter(None, [
        cv.personal.email,
        cv.personal.phone,
        cv.personal.location,
        cv.personal.linkedin,
        cv.personal.github,
        cv.personal.portfolio,
    ]))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, contact, size=9, color=GRAY)
    _rule(p2)

    # ── RESUMEN ────────────────────────────────────────────────
    _section(doc, "RESUMEN PROFESIONAL")
    p = doc.add_paragraph()
    _run(p, cv.summary, size=10)
    p.paragraph_format.space_after = Pt(4)

    # ── HABILIDADES ────────────────────────────────────────────
    _section(doc, "HABILIDADES TÉCNICAS")
    p = doc.add_paragraph()
    _run(p, " • ".join(cv.skills.technical), size=10)

    if cv.skills.tools:
        p = doc.add_paragraph()
        _run(p, "Herramientas: ", size=10, bold=True)
        _run(p, " • ".join(cv.skills.tools), size=10)

    if cv.skills.languages:
        p = doc.add_paragraph()
        _run(p, "Idiomas: ", size=10, bold=True)
        _run(p, " • ".join(cv.skills.languages), size=10)

    # ── EXPERIENCIA ────────────────────────────────────────────
    if cv.experience:
        _section(doc, "EXPERIENCIA PROFESIONAL")
        for exp in cv.experience:
            p = doc.add_paragraph()
            _run(p, exp.title, size=11, bold=True)
            _run(p, f"  |  {exp.start_date} – {exp.end_date}", size=10, color=GRAY)
            p.paragraph_format.space_after = Pt(0)

            p2 = doc.add_paragraph()
            _run(p2, f"{exp.company}  •  {exp.location}", size=10, bold=True, color=BLUE)
            p2.paragraph_format.space_after = Pt(2)

            for ach in exp.achievements:
                b = doc.add_paragraph(style="List Bullet")
                _run(b, ach, size=10)
                b.paragraph_format.left_indent  = Cm(0.5)
                b.paragraph_format.space_after   = Pt(1)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── EDUCACIÓN ─────────────────────────────────────────────
    if cv.education:
        _section(doc, "EDUCACIÓN")
        for edu in cv.education:
            p = doc.add_paragraph()
            _run(p, edu.degree, size=11, bold=True)
            _run(p, f"  |  {edu.start_date} – {edu.end_date}", size=10, color=GRAY)
            p.paragraph_format.space_after = Pt(0)

            p2 = doc.add_paragraph()
            _run(p2, f"{edu.institution}  •  {edu.location}", size=10, bold=True, color=BLUE)
            p2.paragraph_format.space_after = Pt(2)

            if edu.relevant_courses:
                p3 = doc.add_paragraph()
                _run(p3, "Cursos relevantes: ", size=10, bold=True)
                _run(p3, ", ".join(edu.relevant_courses), size=10)

    # ── PROYECTOS ─────────────────────────────────────────────
    if cv.projects:
        _section(doc, "PROYECTOS DESTACADOS")
        for proj in cv.projects:
            p = doc.add_paragraph()
            _run(p, proj.name, size=11, bold=True)
            p.paragraph_format.space_after = Pt(1)

            p2 = doc.add_paragraph()
            _run(p2, proj.description, size=10)
            p2.paragraph_format.space_after = Pt(1)

            p3 = doc.add_paragraph()
            _run(p3, "Tech: ", size=9, bold=True, color=GRAY)
            _run(p3, ", ".join(proj.technologies), size=9, color=GRAY)
            if proj.github:
                _run(p3, f"  |  {proj.github}", size=9, color=BLUE)
            p3.paragraph_format.space_after = Pt(4)

    # ── CERTIFICACIONES ───────────────────────────────────────
    if cv.certifications:
        _section(doc, "CERTIFICACIONES")
        for cert in cv.certifications:
            p = doc.add_paragraph()
            _run(p, f"• {cert}", size=10)

    # Guardar — el filename se sanea porque proviene de datos externos
    OUTPUT_DIR.mkdir(exist_ok=True)
    out = OUTPUT_DIR / _nombre_archivo_seguro(filename)
    # Defensa en profundidad: aunque el saneo fallara, confirmamos que no escapa
    if OUTPUT_DIR.resolve() not in out.resolve().parents:
        raise ValueError(f"Ruta de salida insegura: {out}")
    doc.save(out)
    return out
